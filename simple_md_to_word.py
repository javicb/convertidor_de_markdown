#!/usr/bin/env python3
"""
Convertidor mejorado de Markdown a Word (.docx) con soporte completo para diagramas Mermaid
Versión simplificada que usa directamente python-docx
"""

import os
import re
import asyncio
import tempfile
import shutil
from pathlib import Path
from typing import List, Tuple
import markdown
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from bs4 import BeautifulSoup

class SimpleMdToWordConverter:
    def __init__(self):
        self.temp_dir = None
        self.mermaid_counter = 0
        
    def setup_temp_directory(self) -> str:
        """Crear directorio temporal para archivos de trabajo"""
        self.temp_dir = tempfile.mkdtemp(prefix="md_to_word_")
        return self.temp_dir
        
    def cleanup(self):
        """Limpiar archivos temporales"""
        if self.temp_dir and os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)
    
    async def extract_mermaid_diagrams(self, content: str) -> List[Tuple[str, str]]:
        """Extraer diagramas Mermaid del contenido Markdown"""
        mermaid_pattern = r'```mermaid\n(.*?)\n```'
        diagrams = []
        
        matches = re.finditer(mermaid_pattern, content, re.DOTALL)
        for match in matches:
            mermaid_code = match.group(1).strip()
            self.mermaid_counter += 1
            placeholder_id = f"mermaid_diagram_{self.mermaid_counter}"
            diagrams.append((mermaid_code, placeholder_id))
            
        return diagrams
    
    async def render_mermaid_to_image(self, mermaid_code: str, output_path: str) -> bool:
        """Renderizar diagrama Mermaid a imagen usando Playwright con alta resolución"""
        try:
            async with async_playwright() as p:
                browser = await p.chromium.launch()
                page = await browser.new_page()
                
                # Configurar viewport más grande para mejor resolución
                await page.set_viewport_size({"width": 1920, "height": 1080})
                
                html_template = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <script src="https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js"></script>
                    <style>
                        body {{
                            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                            margin: 40px;
                            background: white;
                            display: flex;
                            justify-content: center;
                            align-items: center;
                            min-height: calc(100vh - 80px);
                        }}
                        .mermaid {{
                            display: flex;
                            justify-content: center;
                            align-items: center;
                            width: 100%;
                            min-height: 400px;
                            font-size: 18px;
                            padding: 30px;
                        }}
                        .mermaid svg {{
                            width: 100% !important;
                            min-width: 800px !important;
                            height: auto !important;
                            font-size: 18px !important;
                            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif !important;
                        }}
                        .mermaid svg text {{
                            font-size: 14px !important;
                            font-weight: 500 !important;
                        }}
                        .mermaid svg .node text {{
                            font-size: 16px !important;
                            font-weight: 600 !important;
                        }}
                    </style>
                </head>
                <body>
                    <div class="mermaid">
                        {mermaid_code}
                    </div>
                    <script>
                        mermaid.initialize({{
                            startOnLoad: true,
                            theme: 'default',
                            flowchart: {{
                                useMaxWidth: true,
                                htmlLabels: true,
                                curve: 'basis'
                            }},
                            themeVariables: {{
                                primaryColor: '#3498db',
                                primaryTextColor: '#2c3e50',
                                primaryBorderColor: '#2980b9',
                                lineColor: '#34495e',
                                sectionBkgColor: '#ecf0f1',
                                altSectionBkgColor: '#bdc3c7',
                                gridColor: '#95a5a6',
                                secondaryColor: '#e74c3c',
                                tertiaryColor: '#f39c12'
                            }}
                        }});
                    </script>
                </body>
                </html>
                """
                
                await page.set_content(html_template)
                await page.wait_for_load_state('networkidle')
                await page.wait_for_selector('.mermaid svg', timeout=15000)
                await page.wait_for_timeout(2000)
                
                mermaid_element = await page.query_selector('.mermaid')
                if mermaid_element:
                    await mermaid_element.screenshot(path=output_path, type='png', omit_background=False)
                    await browser.close()
                    return True
                else:
                    await browser.close()
                    return False
                    
        except Exception as e:
            print(f"Error renderizando diagrama Mermaid: {e}")
            return False
    
    async def process_mermaid_diagrams(self, content: str, temp_dir: str) -> Tuple[str, List[Tuple[str, str, str]]]:
        """Procesar todos los diagramas Mermaid y retornar contenido modificado con rutas de imágenes"""
        self.mermaid_counter = 0
        diagrams = await self.extract_mermaid_diagrams(content)
        
        if not diagrams:
            return content, []
        
        print(f"📊 Encontrados {len(diagrams)} diagramas Mermaid")
        
        # Generar imágenes y crear lista con rutas
        processed_diagrams = []
        modified_content = content
        
        for mermaid_code, placeholder_id in diagrams:
            image_path = os.path.join(temp_dir, f"{placeholder_id}.png")
            print(f"🎨 Generando imagen para {placeholder_id}...")
            
            success = await self.render_mermaid_to_image(mermaid_code, image_path)
            if success:
                print(f"   ✅ Imagen generada: {placeholder_id}.png")
                processed_diagrams.append((mermaid_code, placeholder_id, image_path))
                
                # Reemplazar bloque Mermaid con marcador especial
                mermaid_block = f"```mermaid\n{mermaid_code}\n```"
                placeholder = f"__MERMAID_PLACEHOLDER_{placeholder_id}__"
                modified_content = modified_content.replace(mermaid_block, placeholder)
            else:
                print(f"   ❌ Error generando: {placeholder_id}")
        
        return modified_content, processed_diagrams
    
    def create_word_document_direct(self, content: str, diagrams: List[Tuple[str, str, str]], output_path: str) -> bool:
        """Crear documento Word directamente desde Markdown"""
        try:
            print("📝 Creando documento Word...")
            
            # Crear documento nuevo
            doc = Document()
            
            # Configurar estilos básicos
            self.setup_document_styles(doc)
            
            # Dividir contenido en líneas para procesamiento
            lines = content.split('\n')
            i = 0
            
            while i < len(lines):
                line = lines[i].strip()
                
                # Verificar si es un placeholder de diagrama Mermaid
                if line.startswith('__MERMAID_PLACEHOLDER_'):
                    # Extraer ID del placeholder
                    placeholder_id = line.replace('__MERMAID_PLACEHOLDER_', '').replace('__', '')
                    
                    # Buscar el diagrama correspondiente
                    for j, (mermaid_code, diagram_id, image_path) in enumerate(diagrams):
                        if diagram_id == placeholder_id:
                            print(f"   📷 Insertando diagrama {j+1}: {diagram_id}")
                            
                            # Agregar título del diagrama
                            title_para = doc.add_paragraph()
                            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            title_run = title_para.add_run(f"Diagrama {j+1}: Arquitectura del Sistema")
                            title_run.font.size = Pt(12)
                            title_run.font.bold = True
                            title_run.font.color.rgb = RGBColor(44, 62, 80)  # #2c3e50
                            
                            # Agregar imagen
                            if os.path.exists(image_path):
                                img_para = doc.add_paragraph()
                                img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                run = img_para.add_run()
                                run.add_picture(image_path, width=Inches(6.5))
                                
                                # Espacio después de la imagen
                                doc.add_paragraph()
                                print(f"   ✅ Diagrama {j+1} insertado correctamente")
                            break
                    
                    i += 1
                    continue
                
                # Procesar encabezados
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    if level <= 6:
                        text = line.lstrip('#').strip()
                        if text:
                            heading = doc.add_heading(text, level=level)
                            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                
                # Procesar bloques de código
                elif line.startswith('```'):
                    code_lines = []
                    i += 1
                    while i < len(lines) and not lines[i].strip().startswith('```'):
                        code_lines.append(lines[i])
                        i += 1
                    
                    if code_lines:
                        code_para = doc.add_paragraph()
                        code_run = code_para.add_run('\n'.join(code_lines))
                        code_run.font.name = 'Courier New'
                        code_run.font.size = Pt(9)
                        code_run.font.color.rgb = RGBColor(52, 73, 94)  # Color para código
                
                # Procesar listas
                elif line.startswith('- ') or line.startswith('* ') or re.match(r'\d+\. ', line):
                    text = re.sub(r'^[-*]\s+', '', line)
                    text = re.sub(r'^\d+\.\s+', '', text)
                    if text:
                        para = doc.add_paragraph(style='List Bullet')
                        self.add_formatted_text(para, text)
                
                # Procesar tablas (básico)
                elif '|' in line and line.count('|') >= 2:
                    # Recopilar todas las líneas de la tabla
                    table_lines = [line]
                    i += 1
                    while i < len(lines) and '|' in lines[i]:
                        table_lines.append(lines[i].strip())
                        i += 1
                    i -= 1  # Retroceder una línea
                    
                    if len(table_lines) > 1:
                        self.create_table(doc, table_lines)
                
                # Procesar párrafos normales y listas
                elif line and not line.startswith(('__MERMAID_', '#', '```', '|')) and not re.match(r'^-+\s*$', line) and not re.match(r'\d+\.', line):
                    # Procesar texto con formato markdown básico
                    if line.strip():
                        para = doc.add_paragraph()
                        self.add_formatted_text(para, line)
                
                i += 1
            
            # Guardar documento
            doc.save(output_path)
            print("✅ Documento Word creado exitosamente")
            return True
            
        except Exception as e:
            print(f"❌ Error creando documento Word: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def setup_document_styles(self, doc: Document):
        """Configurar estilos profesionales para el documento"""
        try:
            # Estilo normal
            normal_style = doc.styles['Normal']
            normal_style.font.name = 'Calibri'
            normal_style.font.size = Pt(11)
            
            # Configurar encabezados
            for level in range(1, 7):
                style_name = f'Heading {level}'
                if style_name in doc.styles:
                    style = doc.styles[style_name]
                    style.font.name = 'Calibri'
                    style.font.color.rgb = RGBColor(44, 62, 80)  # #2c3e50
                    
                    sizes = [18, 16, 14, 12, 11, 10]
                    if level <= len(sizes):
                        style.font.size = Pt(sizes[level-1])
                        
        except Exception as e:
            print(f"⚠️  Error configurando estilos: {e}")
    
    def create_table(self, doc: Document, table_lines: List[str]):
        """Crear tabla en Word desde líneas de markdown"""
        try:
            # Procesar líneas de la tabla
            rows = []
            for line in table_lines:
                if '---' not in line:  # Omitir línea separadora
                    cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remover primero y último vacío
                    if cells:
                        rows.append(cells)
            
            if len(rows) < 2:
                return
            
            # Crear tabla
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Light Grid Accent 1'
            
            # Llenar tabla
            for i, row_data in enumerate(rows):
                for j, cell_data in enumerate(row_data):
                    if j < len(table.rows[i].cells):
                        cell = table.rows[i].cells[j]
                        
                        # Limpiar el contenido existente de la celda
                        cell._element.clear()
                        
                        # Crear un párrafo en la celda y aplicar formato Markdown
                        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                        self.add_formatted_text(paragraph, cell_data)
                        
                        # Formato especial para encabezados (primera fila) - aplicar negrita adicional
                        if i == 0:
                            for run in paragraph.runs:
                                run.font.bold = True
            
            doc.add_paragraph()  # Espacio después de la tabla
            
        except Exception as e:
            print(f"⚠️  Error creando tabla: {e}")
    
    def process_markdown_formatting(self, text: str) -> List[Tuple[str, dict]]:
        """Procesar formato de markdown y retornar lista de segmentos con formato"""
        segments = []
        remaining_text = text
        
        # Procesar texto iterativamente para manejar múltiples formatos
        while remaining_text:
            # Buscar el primer patrón de formato
            bold_match = re.search(r'\*\*(.*?)\*\*', remaining_text)
            italic_match = re.search(r'(?<!\*)\*([^*]+?)\*(?!\*)', remaining_text)
            code_match = re.search(r'`(.*?)`', remaining_text)
            
            # Determinar cuál aparece primero
            matches = []
            if bold_match:
                matches.append((bold_match.start(), bold_match.end(), bold_match.group(1), {'bold': True}, bold_match))
            if italic_match:
                matches.append((italic_match.start(), italic_match.end(), italic_match.group(1), {'italic': True}, italic_match))
            if code_match:
                matches.append((code_match.start(), code_match.end(), code_match.group(1), {'code': True}, code_match))
            
            if not matches:
                # No hay más formatos, agregar texto restante
                if remaining_text:
                    segments.append((remaining_text, {}))
                break
            
            # Ordenar por posición y tomar el primero
            matches.sort(key=lambda x: x[0])
            start, end, matched_text, format_dict, match_obj = matches[0]
            
            # Agregar texto normal antes del formato (incluyendo espacios)
            if start > 0:
                normal_text = remaining_text[:start]
                if normal_text:
                    segments.append((normal_text, {}))
            
            # Agregar texto con formato (incluyendo texto vacío si es necesario)
            segments.append((matched_text, format_dict))
            
            # Continuar con el resto del texto
            remaining_text = remaining_text[end:]
        
        # Si no hay segmentos, agregar todo el texto
        if not segments and text:
            segments.append((text, {}))
            
        return segments
    
    def add_formatted_text(self, paragraph, text: str):
        """Agregar texto con formato de markdown al párrafo"""
        segments = self.process_markdown_formatting(text)
        
        for segment_text, format_dict in segments:
            # Crear un run para cada segmento de texto
            run = paragraph.add_run(segment_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            
            # Aplicar formatos explícitamente
            if format_dict.get('bold'):
                run.bold = True
            if format_dict.get('italic'):
                run.italic = True
            if format_dict.get('code'):
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(52, 73, 94)
    
    async def convert_markdown_to_word(self, input_path: str, output_path: str) -> bool:
        """Convertir archivo Markdown a Word"""
        try:
            # Verificar si el archivo de salida existe y está siendo usado
            if os.path.exists(output_path):
                try:
                    # Intentar abrir el archivo para verificar si está en uso
                    with open(output_path, 'r+b') as test_file:
                        pass
                except PermissionError:
                    # El archivo está siendo usado, generar nombre alternativo
                    base_name = os.path.splitext(output_path)[0]
                    ext = os.path.splitext(output_path)[1]
                    counter = 1
                    while os.path.exists(f"{base_name}_{counter}{ext}"):
                        counter += 1
                    output_path = f"{base_name}_{counter}{ext}"
                    print(f"⚠️  Archivo original en uso, generando: {os.path.basename(output_path)}")
            
            temp_dir = self.setup_temp_directory()
            
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            print(f"📖 Procesando: {os.path.basename(input_path)}")
            
            # Procesar diagramas Mermaid
            processed_content, diagrams = await self.process_mermaid_diagrams(content, temp_dir)
            
            # Crear documento Word
            success = self.create_word_document_direct(processed_content, diagrams, output_path)
            
            if success:
                file_size = os.path.getsize(output_path) / (1024 * 1024)
                print(f"✅ Word generado: {os.path.basename(output_path)} ({file_size:.2f} MB)")
            
            return success
            
        except Exception as e:
            print(f"❌ Error en conversión: {e}")
            return False
        finally:
            self.cleanup()

def test_markdown_formatting():
    """Función para probar el procesamiento de formato markdown"""
    print("\n🧪 PROBANDO FORMATO MARKDOWN")
    print("="*30)
    
    converter = SimpleMdToWordConverter()
    
    test_texts = [
        "**Frontend**: Vue 3 + Vite + TypeScript + RDS Energy Components",
        "**Backend**: .NET 10 Web API con Clean Architecture",
        "**Autenticación**: JWT en lugar del sistema actual basado en cookies",
        "**Comunicación**: APIs REST estándar",
        "Este texto tiene **palabras en negrita** y texto normal",
        "Múltiples **formatos** con *cursiva* y `código`"
    ]
    
    for text in test_texts:
        print(f"\nTexto original: {text}")
        segments = converter.process_markdown_formatting(text)
        print("Segmentos procesados:")
        for i, (segment_text, format_dict) in enumerate(segments):
            formats = ", ".join([f"{k}: {v}" for k, v in format_dict.items() if v])
            formats_str = f" [{formats}]" if formats else ""
            print(f"  {i+1}: '{segment_text}'{formats_str}")

async def main():
    """Función principal para probar la conversión"""
    import sys
    
    # Si se llama con --test, ejecutar pruebas
    if len(sys.argv) == 2 and sys.argv[1] == "--test":
        test_markdown_formatting()
        return
    
    if len(sys.argv) != 2:
        print("Uso: python simple_md_to_word.py <archivo.md>")
        print("      python simple_md_to_word.py --test  (para probar formato)")
        return
    
    converter = SimpleMdToWordConverter()
    
    input_file = sys.argv[1]
    output_file = input_file.replace('.md', '.docx')
    
    print("🔄 CONVERTIDOR MD → WORD (Versión Directa)")
    print("="*45)
    
    if os.path.exists(input_file):
        print(f"🚀 Iniciando conversión de {input_file}")
        success = await converter.convert_markdown_to_word(input_file, output_file)
        
        if success:
            print(f"🎉 ¡Conversión a Word completada exitosamente!")
        else:
            print(f"❌ Error en la conversión")
    else:
        print(f"❌ Archivo no encontrado: {input_file}")

if __name__ == "__main__":
    asyncio.run(main())